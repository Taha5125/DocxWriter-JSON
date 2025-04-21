"""Document creator module for DocxWriter."""

import json
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docxwriter.utils.logger import logger
from docxwriter.utils.paths import ensure_dir


class DocumentCreator:
    """Main class for document creation and management."""

    def __init__(
        self, output_dir: Path = Path("data"), watermark_text: Optional[str] = None
    ):
        """Initialize the document creator.

        Args:
            output_dir: Directory where documents will be saved
            watermark_text: Optional watermark text to add to documents
        """
        self.doc = Document()
        self.output_dir = ensure_dir(output_dir)
        self.watermark_text = watermark_text

        # Load styles and fonts from the styles module
        from docxwriter.components import styles

        self.colors = styles.COLORS
        self.fonts = styles.FONTS
        self.sizes = styles.SIZES
        self.line_spacing = styles.LINE_SPACING

        # Apply styles to document
        styles.apply_styles_to_document(self.doc)

    def add_page_break(self) -> None:
        """Add a page break to the document."""
        self.doc.add_page_break()

    def add_image(
        self, image_path: str, width: float = 6.0, height: float = 4.0
    ) -> None:
        """Add an image to the document with specified dimensions."""
        try:
            self.doc.add_picture(image_path, width=Inches(width), height=Inches(height))
        except Exception as e:
            logger.error(f"Error adding image {image_path}: {str(e)}")
            raise

    def add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """Add a hyperlink to the document."""
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create a w:r element and a new w:rPr element
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Join all the xml elements together and add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run()
        r._r.append(hyperlink)

    def set_cell_border(self, cell, **kwargs):
        """
        Set cell border properties.
        Usage:
        set_cell_border(cell, top="single", bottom="single", left="single", right="single")
        """
        # Create border XML elements directly without namespace handling
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Create tcBorders element
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

        # Add borders
        for key, value in kwargs.items():
            if key in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{key}")
                border.set("w:val", value)
                border.set("w:sz", "4")
                border.set("w:space", "0")
                border.set("w:color", "auto")
                tcBorders.append(border)

    def add_table(
        self, data: List[List[str]], style: str = "Table Grid", header_rows: int = 1
    ) -> None:
        """Add a table to the document with the given data."""
        if not data:
            return

        rows = len(data)
        cols = len(data[0])
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style  # Apply built-in style which includes borders

        for i, row in enumerate(data):
            for j, cell_text in enumerate(row):
                if j < len(row):  # Ensure we don't exceed column bounds
                    cell = table.cell(i, j)
                    cell.text = cell_text

                    # Format header rows
                    if i < header_rows:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)

    def add_list(
        self, items: List[str], list_type: str = "bullet", level: int = 0
    ) -> None:
        """Add a bulleted or numbered list to the document."""
        if list_type == "bullet":
            for item in items:
                paragraph = self.doc.add_paragraph(item, style="List Bullet")
                paragraph.paragraph_format.left_indent = Pt(36 * (level + 1))
        else:  # numbered
            for item in items:
                paragraph = self.doc.add_paragraph(item, style="List Number")
                paragraph.paragraph_format.left_indent = Pt(36 * (level + 1))

    def add_section(
        self, heading: str, text: str, level: int = 1, style: str = None
    ) -> None:
        """
        Add a section with heading and formatted paragraphs to the document.

        Args:
            heading: The section heading text
            text: The body text content
            level: Heading level (1-9)
            style: Optional custom style to apply
        """
        try:
            # Add heading
            heading_paragraph = self.doc.add_heading(heading.strip(), level=level)
            heading_run = heading_paragraph.runs[0]
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.name = self.fonts["heading"]

            # Process text content
            paragraphs = text.strip().split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Check if paragraph contains list items
                if para.startswith("- ") or para.startswith("* "):
                    items = [
                        item.strip("- *").strip()
                        for item in para.split("\n")
                        if item.strip()
                    ]
                    self.add_list(items)
                else:
                    paragraph = self.doc.add_paragraph(para)
                    paragraph_run = paragraph.runs[0]
                    paragraph_run.font.size = self.sizes["normal"]
                    paragraph_run.font.name = self.fonts["serif"]

                    if style:
                        paragraph.style = style

        except Exception as e:
            logger.error(f"Error adding section '{heading}': {str(e)}")
            raise

    def process_content(self, content: Dict[str, Any]) -> None:
        """Process the content dictionary and add it to the document."""
        try:
            for heading, text in content.items():
                if isinstance(text, dict):
                    # Handle nested content
                    if "table" in text:
                        self.add_table(
                            text["table"],
                            style=text.get("style", "Table Grid"),
                            header_rows=text.get("header_rows", 1),
                        )
                    elif "list" in text:
                        self.add_list(
                            text["list"],
                            list_type=text.get("list_type", "bullet"),
                            level=text.get("level", 0),
                        )
                    elif "image" in text:
                        self.add_image(
                            text["image"],
                            width=text.get("width", 6.0),
                            height=text.get("height", 4.0),
                        )
                    elif "page_break" in text:
                        self.add_page_break()
                    else:
                        self.add_section(heading, str(text))
                else:
                    self.add_section(heading, str(text))
        except Exception as e:
            logger.error(f"Error processing content: {str(e)}")
            raise

    def add_watermark(self) -> None:
        """Add a hidden watermark to the document if watermark_text is set."""
        if not self.watermark_text:
            return

        watermark = self.doc.add_paragraph(
            self.watermark_text,
            style="Hidden",
        )
        watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark.runs[0].font.size = self.sizes["normal"]
        watermark.runs[0].font.name = self.fonts["heading"]
        watermark.runs[0].font.italic = True
        watermark.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        watermark.runs[0].font.bold = True
        watermark.runs[0].font.underline = True

    def create_document_from_json(self, json_file: str) -> Path:
        """Create a document from a JSON file.

        Args:
            json_file: Path to the JSON file

        Returns:
            Path to the created document
        """
        try:
            # Load and parse JSON data
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Validate required fields
            required_fields = ["title", "file_name", "content"]
            for field in required_fields:
                if field not in data:
                    raise ValueError(f"Missing required field: {field}")

            title = data["title"]
            file_name = data["file_name"]
            sections = data["content"]

            # Add document title
            title_para = self.doc.add_heading(title, 0)
            title_run = title_para.runs[0]
            title_run.font.size = self.sizes["title"]
            title_run.font.bold = True
            title_run.font.name = self.fonts["heading"]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Process content
            if isinstance(sections, list):
                for section in sections:
                    self.process_content(section)
            else:
                self.process_content(sections)

            # Add the watermark if specified
            self.add_watermark()

            # Save the Word document
            doc_path = self.output_dir / file_name
            self.doc.save(doc_path)
            logger.info(f"Document saved successfully to {doc_path}")

            return doc_path

        except FileNotFoundError:
            logger.error(f"Error: {json_file} file not found")
            raise
        except json.JSONDecodeError:
            logger.error(f"Error: Invalid JSON format in {json_file}")
            raise
        except Exception as e:
            logger.error(f"An unexpected error occurred: {str(e)}")
            raise
