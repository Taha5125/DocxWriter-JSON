import json
import logging
from pathlib import Path
from typing import Dict, Any, List
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from styles import doc, FONTS, SIZES, WATERMARK_TEXT

# Set up logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Create output directory if it doesn't exist
output_dir = Path("data")
output_dir.mkdir(exist_ok=True)


def add_page_break() -> None:
    """Add a page break to the document."""
    doc.add_page_break()


def add_image(image_path: str, width: float = 6.0, height: float = 4.0) -> None:
    """Add an image to the document with specified dimensions."""
    try:
        doc.add_picture(image_path, width=Inches(width), height=Inches(height))
    except Exception as e:
        logger.error(f"Error adding image {image_path}: {str(e)}")
        raise


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a hyperlink to the document."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Join all the xml elements together and add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Usage:
    set_cell_border(cell, top="single", bottom="single", left="single", right="single")
    """
    # Create border XML elements directly without namespace handling
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Add borders one by one
    if "top" in kwargs:
        top = OxmlElement("w:top")
        top.set("w:val", kwargs["top"])
        top.set("w:sz", "4")
        top.set("w:space", "0")
        top.set("w:color", "auto")
        tcBorders = tcPr.find("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(top)

    if "bottom" in kwargs:
        bottom = OxmlElement("w:bottom")
        bottom.set("w:val", kwargs["bottom"])
        bottom.set("w:sz", "4")
        bottom.set("w:space", "0")
        bottom.set("w:color", "auto")
        tcBorders = tcPr.find("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(bottom)

    if "left" in kwargs:
        left = OxmlElement("w:left")
        left.set("w:val", kwargs["left"])
        left.set("w:sz", "4")
        left.set("w:space", "0")
        left.set("w:color", "auto")
        tcBorders = tcPr.find("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(left)

    if "right" in kwargs:
        right = OxmlElement("w:right")
        right.set("w:val", kwargs["right"])
        right.set("w:sz", "4")
        right.set("w:space", "0")
        right.set("w:color", "auto")
        tcBorders = tcPr.find("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(right)


def add_table(
    data: List[List[str]], style: str = "Table Grid", header_rows: int = 1
) -> None:
    """Add a table to the document with the given data."""
    if not data:
        return

    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style  # Apply built-in style which includes borders

    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            if j < len(row):  # Ensure we don't exceed column bounds
                cell = table.cell(i, j)
                cell.text = cell_text

                # Format header rows
                if i < header_rows:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)


def add_list(items: List[str], list_type: str = "bullet", level: int = 0) -> None:
    """Add a bulleted or numbered list to the document."""
    if list_type == "bullet":
        for item in items:
            paragraph = doc.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.left_indent = Pt(36 * (level + 1))
    else:  # numbered
        for item in items:
            paragraph = doc.add_paragraph(item, style="List Number")
            paragraph.paragraph_format.left_indent = Pt(36 * (level + 1))


def add_section(heading: str, text: str, level: int = 1, style: str = None) -> None:
    """
    Add a section with heading and formatted paragraphs to the document.

    Args:
        heading: The section heading text
        text: The body text content
        level: Heading level (1-9)
        style: Optional custom style to apply
    """
    try:
        # Add heading
        heading_paragraph = doc.add_heading(heading.strip(), level=level)
        heading_run = heading_paragraph.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.name = FONTS["heading"]

        # Process text content
        paragraphs = text.strip().split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if paragraph contains list items
            if para.startswith("- ") or para.startswith("* "):
                items = [
                    item.strip("- *").strip()
                    for item in para.split("\n")
                    if item.strip()
                ]
                add_list(items)
            else:
                paragraph = doc.add_paragraph(para)
                paragraph_run = paragraph.runs[0]
                paragraph_run.font.size = SIZES["normal"]
                paragraph_run.font.name = FONTS["serif"]

                if style:
                    paragraph.style = style

    except Exception as e:
        logger.error(f"Error adding section '{heading}': {str(e)}")
        raise


def process_content(content: Dict[str, Any]) -> None:
    """Process the content dictionary and add it to the document."""
    try:
        for heading, text in content.items():
            if isinstance(text, dict):
                # Handle nested content
                if "table" in text:
                    add_table(
                        text["table"],
                        style=text.get("style", "Table Grid"),
                        header_rows=text.get("header_rows", 1),
                    )
                elif "list" in text:
                    add_list(
                        text["list"],
                        list_type=text.get("list_type", "bullet"),
                        level=text.get("level", 0),
                    )
                elif "image" in text:
                    add_image(
                        text["image"],
                        width=text.get("width", 6.0),
                        height=text.get("height", 4.0),
                    )
                elif "page_break" in text:
                    add_page_break()
                else:
                    add_section(heading, str(text))
            else:
                add_section(heading, str(text))
    except Exception as e:
        logger.error(f"Error processing content: {str(e)}")
        raise


def main():
    """Main function to process the JSON data and create the document
    Data Structure:
        {
            "title": "Document Title",
            "file_name": "document.docx",
            "content": [
                {"heading": "Section Heading", "text": "Section content"}
            ]
        }
    """
    try:
        # Load and parse JSON data
        with open("data.json", "r", encoding="utf-8") as f:
            data = json.load(f)

        # Validate required fields
        required_fields = ["title", "file_name", "content"]
        for field in required_fields:
            if field not in data:
                raise ValueError(f"Missing required field: {field}")

        title = data["title"]
        file_name = data["file_name"]
        sections = data["content"]

        # Add document title
        title_para = doc.add_heading(title, 0)
        title_run = title_para.runs[0]
        title_run.font.size = SIZES["title"]
        title_run.font.bold = True
        title_run.font.name = FONTS["heading"]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process content
        if isinstance(sections, list):
            for section in sections:
                process_content(section)
        else:
            process_content(sections)

        ## Adding a hiddin watermark
        if WATERMARK_TEXT:
            watermark = doc.add_paragraph(
                WATERMARK_TEXT,
                style="Hidden",
            )
            watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
            watermark.runs[0].font.size = SIZES["normal"]
            watermark.runs[0].font.name = FONTS["heading"]
            watermark.runs[0].font.italic = True
            watermark.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            watermark.runs[0].font.bold = True
            watermark.runs[0].font.underline = True

        # Save the Word document
        doc_path = output_dir / file_name
        doc.save(doc_path)
        logger.info(f"Document saved successfully to {doc_path}")

    except FileNotFoundError:
        logger.error("Error: data.json file not found")
        raise
    except json.JSONDecodeError:
        logger.error("Error: Invalid JSON format in data.json")
        raise
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        raise


if __name__ == "__main__":
    main()
