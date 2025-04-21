"""Tests for the DocxWriter project."""

import os
import json
import unittest
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Import the modules to test - use the compatibility layer for maximum test stability
import writer
import styles
from docxwriter.utils.paths import output_dir


class TestDocxWriter(unittest.TestCase):
    """Test suite for the DocxWriter project."""

    def setUp(self):
        """Set up test environment before each test."""
        # Create a temporary directory for test outputs
        self.temp_dir = tempfile.TemporaryDirectory()
        self.output_dir = Path(self.temp_dir.name)

        # Create a sample test_data.json file
        self.sample_data = {
            "title": "Test Document",
            "file_name": "test_output.docx",
            "content": {
                "Introduction": "This is a test introduction.",
                "Table Section": {
                    "table": [["Header 1", "Header 2"], ["Data 1", "Data 2"]]
                },
                "List Section": {
                    "list": ["Item 1", "Item 2", "Item 3"],
                    "list_type": "bullet",
                },
            },
        }

        with open(self.output_dir / "test_data.json", "w", encoding="utf-8") as f:
            json.dump(self.sample_data, f)

        # Cache the original output directory
        self.original_output_dir = output_dir

        # Create a sample image for testing
        self.image_path = self.output_dir / "test_image.png"
        with open(self.image_path, "w") as f:
            f.write("dummy image data")

        # Use document from writer (via compat layer)
        self.doc = writer.doc

    def tearDown(self):
        """Clean up after each test."""
        # Clean up temp directory
        self.temp_dir.cleanup()

    def test_add_page_break(self):
        """Test adding a page break to the document."""
        # Add some content
        self.doc.add_paragraph("Before page break")

        # Add a page break
        writer.add_page_break()

        # Add more content
        self.doc.add_paragraph("After page break")

        # Save the document
        output_path = self.output_dir / "page_break_test.docx"
        self.doc.save(output_path)

        # Verify the document was created
        self.assertTrue(output_path.exists())

    def test_add_table(self):
        """Test adding a table to the document."""
        # Sample table data
        data = [
            ["Header 1", "Header 2"],
            ["Row 1 Col 1", "Row 1 Col 2"],
            ["Row 2 Col 1", "Row 2 Col 2"],
        ]

        # Add the table
        writer.add_table(data, style="Table Grid", header_rows=1)

        # Save the document
        output_path = self.output_dir / "table_test.docx"
        self.doc.save(output_path)

        # Verify the document was created
        self.assertTrue(output_path.exists())

        # Print success message
        print("Table test document saved successfully")

    def test_add_list(self):
        """Test adding a list to the document."""
        # Sample list items
        items = ["Item 1", "Item 2", "Item 3"]

        # Add a bulleted list
        writer.add_list(items, list_type="bullet")

        # Save the document
        output_path = self.output_dir / "list_test.docx"
        self.doc.save(output_path)

        # Verify the document was created
        self.assertTrue(output_path.exists())

        # Print success message
        print("List test document saved successfully")

    def test_add_section(self):
        """Test adding a section to the document."""
        # Add a section with specific content
        writer.add_section(
            "Test Heading",
            "This is a test section with multiple paragraphs.\n\nThis is the second paragraph.",
        )

        # Save the document
        output_path = self.output_dir / "section_test.docx"
        self.doc.save(output_path)

        # Verify the document was created
        self.assertTrue(output_path.exists())

        # Print success message
        print("Section test document saved successfully")

    def test_process_content(self):
        """Test processing content from a dictionary."""
        # Sample content
        content = {"Section 1": "This is section 1."}

        # Process the content
        writer.process_content(content)

        # Save the document
        output_path = self.output_dir / "process_content_test.docx"
        self.doc.save(output_path)

        # Verify the document was created
        self.assertTrue(output_path.exists())

        # Print success message
        print("Process content test document saved successfully")

    def test_main_function(self):
        """Test the main function with a sample data.json file."""
        # Temporarily change the output directory
        writer.output_dir = self.output_dir

        # Save current directory
        current_dir = os.getcwd()

        try:
            # Create a test data.json file in the current directory
            with open("data.json", "w", encoding="utf-8") as f:
                json.dump(self.sample_data, f)

            # Run the main function
            writer.main()

            # Verify the output file was created
            output_path = self.output_dir / self.sample_data["file_name"]

            # Print debug info
            print(f"Output path: {output_path}")
            print(f"Output path exists: {output_path.exists()}")
            print(f"Output directory contents: {list(self.output_dir.glob('*'))}")

            # Check if file exists
            self.assertTrue(
                output_path.exists(), f"Output file was not created at {output_path}"
            )

        finally:
            # Clean up test data file
            if os.path.exists("data.json"):
                os.remove("data.json")

            # Restore output directory
            writer.output_dir = self.original_output_dir

    def test_styles(self):
        """Test the styles module."""
        # Test color definitions
        self.assertIsInstance(styles.COLORS["primary"], RGBColor)
        self.assertIsInstance(styles.COLORS["error"], RGBColor)

        # Test font definitions
        self.assertIsInstance(styles.FONTS["serif"], str)
        self.assertIsInstance(styles.FONTS["heading"], str)

        # Test size definitions
        self.assertIsInstance(styles.SIZES["normal"], Pt)
        self.assertIsInstance(styles.SIZES["title"], Pt)

        # Test line spacing definitions
        self.assertIsInstance(styles.LINE_SPACING["single"], float)
        self.assertIsInstance(styles.LINE_SPACING["double"], float)

        # Test style creation - provide doc as first parameter
        style = styles.create_style(styles.doc, "TestStyle")
        self.assertIsNotNone(style)

        # Test paragraph formatting
        styles.apply_paragraph_format(style, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(style.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)

        # Test font formatting
        styles.apply_font_format(style, bold=True, size=Pt(14))
        self.assertTrue(style.font.bold)
        self.assertEqual(style.font.size, Pt(14))


if __name__ == "__main__":
    unittest.main()
